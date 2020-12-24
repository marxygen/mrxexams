import docx, os, json, random
from datetime import datetime as dt
from params import FILES_PATH

def create_document(filename, categories):
    try:
        with open(os.path.join(FILES_PATH, filename.lower()), 'rb') as f:
            js = list(json.load(f))
            questions = []
            for element in js:
                category = element['category']
                if category.lower() not in categories:
                    continue
                for item in element['items']:
                    text = item[0]
                    answer = item[1]
                    attachments = item[2]

                    questions.append([category, text, answer, attachments])

            doc = docx.Document()
            categories_str = ','.join([category.capitalize() for category in list(categories)])
            id = ''.join(([str(random.randint(0,10)) for _ in range(0, 3)] + [str(dt.now())[:10]]))
            doc.add_heading(f'{id}\nQuestions in categories {categories_str}', 2)
            doc.add_paragraph(f'This test contains {len(questions)} questions')
            
            answer = docx.Document()
            answer.add_heading(f'Answers for test with id {id} in categories {categories_str}', 2)
            answer.add_paragraph(f'This test contains {len(questions)} questions')

            q_par = doc.add_paragraph('')
            ans_par = answer.add_paragraph('')

            for q_num in range(1, len(questions)+1):
                question = random.choice(questions)
               
                q_par.add_run(f"{q_num}) {question[1]} ({question[0]})?\n")
                q_par.add_run('_'*70 + '\n\n')
                if len(question[3]) > 2:
                    for pic in eval(question[3]):
                        try:
                            doc.add_picture(os.path.join(FILES_PATH, pic))
                        except FileNotFoundError:
                            continue
                ans_par.add_run(f"{q_num}) {question[2]} ({question[0]})\n")

                questions.remove(question)

            test_name = os.path.join(FILES_PATH, 'tests', f'(ID{id}) Test dd {str(dt.now())[:10]}.docx')
            doc.save(test_name)
            answers_file_name = os.path.join(FILES_PATH, 'tests', f'ANSWERS FOR THE TEST {id}.docx')
            answer.save(answers_file_name)
            return (test_name, answers_file_name)
    except Exception as e:
        print(e)
        return None